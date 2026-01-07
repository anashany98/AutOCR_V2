"""
Content extraction utilities for multiple document types.

This module centralises the logic required to derive text content from a wide
range of supported file formats so that downstream OCR/classification flows
can work with heterogeneous inputs (images, PDFs, Office documents, emails,
plain text files, etc.).
"""

from __future__ import annotations

import json
import logging
import os
from email import policy
from email.parser import BytesParser
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, Optional, Tuple

try:
    import docx  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    docx = None  # type: ignore

try:
    import openpyxl  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    openpyxl = None  # type: ignore

try:
    import extract_msg  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    extract_msg = None  # type: ignore

from .ocr_manager import OCRManager

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".gif", ".webp", ".jfif", ".avif"}
PDF_EXTENSIONS = {".pdf"}
TEXT_EXTENSIONS = {".txt", ".csv", ".json", ".md", ".log"}
DOC_EXTENSIONS = {".docx"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm"}
EMAIL_EXTENSIONS = {".eml"}
MSG_EXTENSIONS = {".msg"}


class _HTMLToTextParser(HTMLParser):
    """Basic HTML to text converter using the standard library."""

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []

    def handle_data(self, data: str) -> None:  # pragma: no cover - trivial
        if data:
            self._chunks.append(data)

    def handle_entityref(self, name: str) -> None:  # pragma: no cover - trivial
        self._chunks.append(self.unescape(f"&{name};"))

    def handle_charref(self, name: str) -> None:  # pragma: no cover - trivial
        self._chunks.append(self.unescape(f"&#{name};"))

    def get_text(self) -> str:
        return " ".join(chunk.strip() for chunk in self._chunks if chunk.strip())


def _html_to_text(html: str) -> str:
    parser = _HTMLToTextParser()
    parser.feed(html)
    return parser.get_text()


def _read_text_file(path: Path, encoding_candidates: Iterable[str]) -> str:
    for encoding in encoding_candidates:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_email_text(path: Path) -> str:
    with path.open("rb") as fh:
        message = BytesParser(policy=policy.default).parse(fh)

    parts: list[str] = []
    for header in ("Subject", "From", "To", "Cc"):
        value = message.get(header)
        if value:
            parts.append(f"{header}: {value}")

    for payload in message.walk():
        if payload.is_multipart():
            continue
        content_type = payload.get_content_type()
        try:
            body = payload.get_content()
        except Exception:
            body = payload.get_payload(decode=True) or b""
            body = body.decode("utf-8", errors="replace")

        if content_type == "text/plain":
            parts.append(body)
        elif content_type == "text/html":
            parts.append(_html_to_text(body))

    return "\n".join(part.strip() for part in parts if part and str(part).strip())


def _extract_docx_text(path: Path) -> str:
    if docx is None:
        raise RuntimeError(
            "python-docx is required to extract DOCX content but is not installed"
        )

    document = docx.Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    # Include table content.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


def _extract_excel_text(path: Path) -> str:
    if openpyxl is None:
        raise RuntimeError(
            "openpyxl is required to extract Excel content but is not installed"
        )

    workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"[Hoja: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell not in (None, "")]
            if row_values:
                parts.append(" | ".join(row_values))
    workbook.close()
    return "\n".join(parts)


def _extract_msg_text(path: Path) -> str:
    if extract_msg is None:
        raise RuntimeError(
            "extract-msg is required to extract MSG content but is not installed"
        )

    msg = extract_msg.Message(str(path))
    parts: list[str] = []
    
    # Headers
    if msg.subject:
        parts.append(f"Subject: {msg.subject}")
    if msg.sender:
        parts.append(f"From: {msg.sender}")
    if msg.to:
        parts.append(f"To: {msg.to}")
    if msg.date:
        parts.append(f"Date: {msg.date}")
    
    # Body
    if msg.body:
        parts.append("\nBody:")
        parts.append(msg.body)
        
    msg.close()
    return "\n".join(part.strip() for part in parts if part and str(part).strip())


def extract_content(
    file_path: str,
    ocr_engine: Optional[OCRManager] = None,
    logger: Optional[logging.Logger] = None,
) -> Tuple[str, Optional[str], float, bool]:
    """
    Extract textual content from ``file_path``.

    Parameters
    ----------
    file_path:
        Absolute path to the document to analyse.
    ocr_engine:
        OCR manager used for image-based extraction when required.
    logger:
        Optional logger for diagnostics.

    Returns
    -------
    tuple[str, Optional[str]]
        Extracted text and an optional language identifier.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    try:
        if ext in IMAGE_EXTENSIONS or ext in PDF_EXTENSIONS:
            if not ocr_engine:
                raise RuntimeError("OCR engine is not available for image-based files")
            text, language, confidence, is_handwritten = ocr_engine.extract_text(file_path)
            return text, language, confidence, is_handwritten

        if ext in TEXT_EXTENSIONS:
            text = _read_text_file(path, ("utf-8", "latin-1", "cp1252"))
            return text, None, 1.0, False

        if ext in EMAIL_EXTENSIONS:
            return _extract_email_text(path), None, 1.0, False

        if ext in DOC_EXTENSIONS:
            return _extract_docx_text(path), None, 1.0, False

        if ext in EXCEL_EXTENSIONS:
            return _extract_excel_text(path), None, 1.0, False

        if ext in MSG_EXTENSIONS:
            return _extract_msg_text(path), None, 1.0, False

        if ext == ".json":
            # Already handled by TEXT_EXTENSIONS, but keep explicit branch.
            return _read_text_file(path, ("utf-8",)), None, 1.0, False

        # Unknown formats fall back to binary read and best-effort decoding.
        text = _read_text_file(path, ("utf-8", "latin-1", "cp1252"))
        return text, None, 0.2, False
    except Exception as exc:
        if logger:
            logger.error(f"Failed to extract content from {path.name}: {exc}")
        raise
